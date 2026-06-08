"""
gen_test_doc.py
───────────────
Sinh tài liệu Word "Hướng dẫn test phần mềm — Điều khiển trực tiếp robot thật
(Phase 1 & 2)" từ nội dung soạn sẵn (an toàn + trình tự test theo bậc).

    .venv\\Scripts\\python.exe scripts\\gen_test_doc.py
Kết quả: docs/HUONG_DAN_TEST_DIEU_KHIEN.docx
"""
from __future__ import annotations

import sys
from pathlib import Path

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:                       # noqa: BLE001
    pass

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "HUONG_DAN_TEST_DIEU_KHIEN.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
WARN = RGBColor(0xC0, 0x00, 0x00)


def main() -> int:
    doc = Document()
    base = doc.styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(11)

    def h(text, level=1):
        p = doc.add_heading(text, level=level)
        for r in p.runs:
            r.font.color.rgb = ACCENT
        return p

    def para(text="", bold=False, color=None, italic=False, size=None):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        if color is not None:
            r.font.color.rgb = color
        if size is not None:
            r.font.size = Pt(size)
        return p

    def bullet(text, style="List Bullet"):
        doc.add_paragraph(text, style=style)

    def check(text):
        doc.add_paragraph(f"☐  {text}")        # ☐ ô tick

    def table(headers, rows, widths=None):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, htext in enumerate(headers):
            c = t.rows[0].cells[i]
            c.text = ""
            run = c.paragraphs[0].add_run(htext)
            run.bold = True
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
        return t

    # ─────────────────── Trang tiêu đề ───────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("HƯỚNG DẪN TEST PHẦN MỀM")
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = ACCENT
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Điều khiển trực tiếp robot thật qua HSE — Phase 1 (Send pose) "
                    "& Phase 2 (Live jog)")
    r.bold = True; r.font.size = Pt(13)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Dự án: DTwinGP7 — Digital Twin Yaskawa GP7 / YRC1000\n"
                 "Phạm vi: kiểm thử chức năng điều khiển trực tiếp (không nạp job)\n"
                 "Phiên bản tài liệu: 1.0").italic = True

    para()
    box = para()
    rr = box.add_run("⚠ CẢNH BÁO: Đây là chuyển động THẬT. Mọi bậc test phải có "
                     "workspace trống, tay luôn đặt trên E-stop vật lý, YRC1000 ở "
                     "REMOTE mode. Streaming (Phase 2) CHƯA được kiểm chứng phần "
                     "cứng — bắt đầu với jog nhỏ và chậm.")
    rr.bold = True; rr.font.color.rgb = WARN

    doc.add_page_break()

    # ─────────────────── 1. Mục đích ───────────────────
    h("1. Mục đích & phạm vi", 1)
    para("Tài liệu hướng dẫn kiểm thử hai chức năng điều khiển trực tiếp robot "
         "Yaskawa GP7 qua giao thức HSE (không nạp file .JBI):")
    bullet("Phase 1 — Send pose: gửi pose hiện tại xuống robot thật, rời rạc (1 lần/lần bấm).")
    bullet("Phase 2 — Live jog: stream mọi thao tác jog xuống robot thật ~8 Hz (kiểu RoboDK).")
    para("Mục tiêu: xác nhận robot di chuyển ĐÚNG (hướng, vị trí, tốc độ) và AN "
         "TOÀN (dừng được, không vượt giới hạn, không nhảy đột ngột) trước khi đưa "
         "vào sử dụng.")

    # ─────────────────── 2. An toàn ───────────────────
    h("2. Quy tắc an toàn bắt buộc", 1)
    bullet("Workspace trống quanh robot, không người trong tầm với.")
    bullet("Tay luôn đặt trên E-stop vật lý (TP hoặc tủ điện) trong suốt quá trình test.")
    bullet("YRC1000 ở REMOTE mode; HSE Server bật; TOOL đã setup; ping OK.")
    bullet("Tốc độ giới hạn ≤ 10% (cap cứng trong phần mềm).")
    bullet("Mỗi bậc chỉ thực hiện khi bậc trước đã PASS; tăng biên độ/tốc độ từ từ.")
    bullet("Bất kỳ chuyển động/âm thanh ngoài dự kiến: bấm E-stop vật lý TRƯỚC, phân tích SAU.")
    para("Lưu ý phần mềm: nút Stop / tắt toggle / đóng app đều gửi lệnh servo-OFF, "
         "NHƯNG không bao giờ thay thế E-stop vật lý.", italic=True)

    # ─────────────────── 3. Chuẩn bị ───────────────────
    h("3. Chuẩn bị trước khi test", 1)
    check("Workspace trống, E-stop trong tầm tay.")
    check("YRC1000: REMOTE mode, HSE Server bật, TOOL đã setup, ping tới IP OK.")
    check("Hạ tốc độ cap: Program → Post-processor settings… → đặt max_speed_pct = 3 "
          "(live-jog/send-pose dùng min(max_speed_pct, 10) → chạy 3%).")
    check("Mở app → Load Robot GP7.")
    check("Robot → Connection settings… → nhập IP → bấm Test (phải thấy joints + không alarm).")

    # ─────────────────── 4. Rủi ro cần kiểm chứng ───────────────────
    h("4. Các rủi ro cần kiểm chứng (đối tượng của test)", 1)
    table(
        ["#", "Mức", "Rủi ro", "Bậc kiểm"],
        [
            ["1", "CAO NHẤT", "Stream lệnh MOVE khi robot đang chạy dở → controller "
             "có thể xếp hàng lệnh → robot chạy tiếp/quá đà sau khi ngừng jog", "Bậc 3"],
            ["2", "Cao", "Quy ước pulse khi GHI (MOVE) khác khi ĐỌC → trục (đặc biệt U) "
             "có thể sai vị trí", "Bậc 1, 2"],
            ["3", "Cao", "Đơn vị speed của lệnh 0x8B hiểu sai → robot chạy nhanh hơn 10%",
             "Bậc 2"],
            ["4", "TB", "Stop() (servo-OFF) chưa chứng minh dừng được robot đang chạy",
             "Bậc 2"],
        ],
    )

    # ─────────────────── 5. Trình tự test ───────────────────
    h("5. Trình tự test theo bậc", 1)

    h("Bậc 0 — Offline (chưa kết nối / rút mạng)", 2)
    para("Mục đích: kiểm luồng phần mềm, KHÔNG có chuyển động.")
    table(["Thao tác", "Kỳ vọng (PASS)"],
          [["Bật Robot → ◀▶ Live jog khi không kết nối được",
            "Hiện dialog → Yes → báo \"HSE not responding — live jog OFF\", "
            "dấu tích tự mất"],
           ["Bật khi đang chạy Mirror/experiment",
            "Báo \"Robot busy…\", từ chối, dấu tích tự mất"]])
    para("PASS nếu mọi trường hợp lỗi đều tự bỏ tích và không treo.", italic=True)

    h("Bậc 1 — Mirror (chỉ ĐỌC, robot KHÔNG nhận lệnh)", 2)
    para("Mục đích: kiểm chứng quy ước pulse↔độ chiều ĐỌC (rủi ro #2).")
    check("Digital Twin → Start live mirror.")
    check("Trên TP, jog tay TỪNG trục (S, L, U, R, B, T) tới một góc đã biết.")
    check("So giá trị joint trên app với TP (đơn vị độ) từng trục.")
    para("PASS: cả 6 trục khớp (sai < ~0.5°). FAIL: trục nào lệch (đặc biệt U) → "
         "dừng, nghi ghép trục L–U.", italic=True)

    h("Bậc 2 — Send pose (Phase 1, rời rạc)", 2)
    para("Mục đích: kiểm chứng rủi ro #2 (ghi), #3 (speed), #4 (stop). Bắt đầu với "
         "dịch chuyển NHỎ.")
    para("Các bước:")
    bullet("1) Jog robot ẢO lệch ~5° MỘT trục so với pose thực (vd chỉ U +5°).")
    bullet("2) Robot → ⬇ Send current pose to REAL robot → đọc kỹ joints đích → Yes.")
    bullet("3) Quan sát theo bảng dưới.")
    table(["Kiểm", "Kỳ vọng (PASS)", "Nếu sai"],
          [["Hướng", "Robot đi đúng trục/đúng chiều ~5°", "#2 sai → DỪNG"],
           ["Tốc độ", "Chậm (≈3%), không giật nhanh", "#3 sai → E-stop, DỪNG"],
           ["Đến đích", "Pose thực ≈ pose app (so trên TP)", "#2 sai → DỪNG"]])
    para("4) Test STOP (#4): lặp lại với dịch chuyển ~15°; NGAY khi robot bắt đầu "
         "chạy → bấm Program → Stop (■).")
    table(["Kiểm", "Kỳ vọng"],
          [["Stop", "Robot dừng ngay, servo rớt (báo \"servo OFF\")"]])
    para("PASS cả 4 → Phase 1 dùng được.", italic=True)

    h("Bậc 3 — Live jog 1 BƯỚC (Phase 2) — phép thử then chốt", 2)
    para("Mục đích: kiểm chứng rủi ro #1 (overshoot do xếp hàng lệnh).", bold=True)
    bullet("1) Robot → ◀▶ Live jog (ON) → Yes → đợi báo \"LIVE JOG ON… synced to "
           "real pose\" (app tự đồng bộ ảo về thực).")
    bullet("2) Chọn Cartesian jog, step nhỏ (vd 5 mm hoặc 2°), jog ĐÚNG 1 bước rồi "
           "BUÔNG TAY hoàn toàn.")
    bullet("3) Quan sát 5–10 giây.")
    table(["Kiểm", "Kỳ vọng (PASS)", "Nếu sai (FAIL)"],
          [["#1 overshoot", "Robot đi 1 đoạn nhỏ rồi DỪNG hẳn",
            "Robot chạy tiếp/quá đà/giật sau khi buông → controller xếp hàng lệnh → "
            "E-stop ngay, tắt toggle, BÁO LẠI để đổi cơ chế"],
           ["Đồng bộ", "Pose thực ≈ pose ảo", "Lệch → kiểm lại Bậc 1/2"]])
    para("Chỉ qua Bậc 4 nếu robot DỪNG dứt khoát sau mỗi bước rời.", bold=True)

    h("Bậc 4 — Live jog LIÊN TỤC", 2)
    bullet("1) Vẫn ON, jog liên tục CHẬM (giữ dial/joint slider kéo từ từ) — KHÔNG "
           "kéo mạnh/đột ngột.")
    bullet("2) Kỳ vọng: robot bám mượt, dừng khi ngừng jog.")
    bullet("3) Thử teleport-guard: đang ON, bấm Home (hoặc paste target xa) rồi jog → "
           "app phải báo \"jog step …° BLOCKED\", robot KHÔNG nhảy.")
    bullet("4) Tắt toggle → servo OFF.")
    para("PASS toàn bộ → Phase 2 dùng được ở tốc thấp; sau đó mới cân nhắc tăng "
         "max_speed_pct.", italic=True)

    # ─────────────────── 6. Quy tắc dừng ───────────────────
    h("6. Quy tắc dừng chung", 1)
    bullet("Bất kỳ bất thường → E-stop vật lý TRƯỚC.")
    bullet("Mỗi lần lên bậc: tăng biên độ/tốc độ từ từ.")
    bullet("Ghi lại kết quả #1–#4 (bảng mục 7) để chốt phần \"unverified\" hoặc sửa "
           "phần mềm nếu controller hành xử khác.")

    # ─────────────────── 7. Bảng ghi kết quả ───────────────────
    h("7. Bảng ghi kết quả (sign-off)", 1)
    table(["Bậc", "Nội dung", "PASS/FAIL", "Ghi chú", "Người test / ngày"],
          [["0", "Offline — luồng phần mềm", "", "", ""],
           ["1", "Mirror — pulse↔độ (đọc)", "", "", ""],
           ["2", "Send pose — hướng/tốc/đích/stop", "", "", ""],
           ["3", "Live jog 1 bước — overshoot (#1)", "", "", ""],
           ["4", "Live jog liên tục + teleport guard", "", "", ""]])
    para()
    para("Kết luận chung: ☐ ĐẠT   ☐ KHÔNG ĐẠT   — chữ ký: __________________", bold=True)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Đã xuất: {OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
